
import streamlit as st
from docx import Document
import io
import json
import re
import tempfile
import os
from datetime import datetime
from typing import Dict, Any, List, Tuple, Optional

import pdfplumber
from PIL import Image

# Try easyocr first (pure python). If not installed, fall back to pytesseract.
try:
    import easyocr
    EASYOCR_AVAILABLE = True
    _easyocr_reader = None
except Exception:
    EASYOCR_AVAILABLE = False
    _easyocr_reader = None

try:
    import pytesseract
    PYTESSERACT_AVAILABLE = True
except Exception:
    PYTESSERACT_AVAILABLE = False

# pdf2image (only for fallback OCR path if easyocr missing)
try:
    from pdf2image import convert_from_path
    PDF2IMAGE_AVAILABLE = True
except Exception:
    PDF2IMAGE_AVAILABLE = False

# Optional Gemini LLM client will be imported at runtime
try:
    import google.generativeai as genai  # type: ignore
except Exception:
    genai = None

# ---------- Helper functions ----------

def init_easyocr(lang_list: List[str] = ["en"]) -> None:
    global _easyocr_reader
    if EASYOCR_AVAILABLE and _easyocr_reader is None:
        _easyocr_reader = easyocr.Reader(lang_list, gpu=False)

def ocr_image_with_easyocr(img: Image.Image) -> str:
    """Return OCR text using easyocr for a PIL Image."""
    init_easyocr()
    if not _easyocr_reader:
        return ""
    try:
        results = _easyocr_reader.readtext(np.array(img))
        # results is list of (bbox, text, conf)
        return "\n".join([r[1] for r in results])
    except Exception:
        return ""

def ocr_image_with_pytesseract(img: Image.Image) -> str:
    try:
        return pytesseract.image_to_string(img)
    except Exception:
        return ""

def image_to_text(img: Image.Image) -> str:
    """Try easyocr first, then pytesseract. Return text (possibly empty)."""
    text = ""
    if EASYOCR_AVAILABLE:
        try:
            init_easyocr()
            # try/except around easyocr
            results = _easyocr_reader.readtext(np.array(img))
            text = "\n".join([r[1] for r in results])
            if text.strip():
                return text
        except Exception:
            text = ""
    if PYTESSERACT_AVAILABLE:
        try:
            return pytesseract.image_to_string(img)
        except Exception:
            return ""
    return text

# small local import to avoid top-level numpy if not needed
def _import_numpy():
    # import lazily because easyocr uses numpy
    try:
        import numpy as np
        return np
    except Exception:
        return None

# ---------- Text extraction ----------

def extract_pages_from_pdf(path: str) -> List[str]:
    """
    Return a list of page texts. For each page try pdfplumber extraction; if page empty,
    attempt OCR on the page image (using easyocr or pytesseract + pdf2image).
    """
    page_texts: List[str] = []
    try:
        with pdfplumber.open(path) as pdf:
            for i, page in enumerate(pdf.pages):
                try:
                    txt = page.extract_text() or ""
                    if txt and len(txt.strip()) > 30:
                        page_texts.append(txt.replace("\n", " ").strip())
                        continue
                except Exception:
                    txt = ""
                # fallback OCR for that page
                # try cropping/using page.to_image if available, else use pdf2image conversion of that page
                ocr_text = ""
                try:
                    # If pdfplumber supports to_image, use it:
                    try:
                        pil_image = page.to_image(resolution=300).original
                        ocr_text = image_to_text(pil_image)
                    except Exception:
                        # fallback: use pdf2image to convert page i to image
                        if PDF2IMAGE_AVAILABLE:
                            imgs = convert_from_path(path, first_page=i+1, last_page=i+1, dpi=300)
                            if imgs:
                                ocr_text = image_to_text(imgs[0])
                except Exception:
                    ocr_text = ""
                page_texts.append(ocr_text.replace("\n", " ").strip() if ocr_text else "")
    except Exception:
        # If pdfplumber can't open, fallback to pdf2image+OCR for entire doc
        if PDF2IMAGE_AVAILABLE:
            try:
                imgs = convert_from_path(path, dpi=300)
                for img in imgs:
                    page_texts.append(image_to_text(img))
            except Exception:
                pass
    return page_texts

def extract_text_from_image_file(path: str) -> str:
    try:
        img = Image.open(path)
        return image_to_text(img) or ""
    except Exception:
        return ""

# ---------- Field extraction & prompt ----------

# You should customize this mapping to include synonyms that appear in your reports.
# The mapping maps template keys to likely label patterns.
DEFAULT_FIELD_LABEL_MAP = {
    "INSURED_NAME": ["insured", "insured name", "client name", "named insured"],
    "INSURED_H_STREET": ["address", "insured address", "street"],
    "INSURED_H_CITY": ["city"],
    "INSURED_H_STATE": ["state"],
    "INSURED_H_ZIP": ["zip", "postal code", "zipcode"],
    "DATE_INSPECTED": ["date inspected", "inspection date"],
    "DATE_LOSS": ["date of loss", "loss date"],
    "DATE_RECEIVED": ["date received", "received date"],
    "MORTGAGEE": ["mortgagee", "mortgage party"],
    "MORTGAGE_CO": ["mortgage company"],
    "TOL_CODE": ["tol code", "t o l code", "loss code"]
}

def build_strict_prompt_for_field(field_key: str, label_hints: List[str], page_text: str) -> str:
    """
    Builds a strict instruction prompt for the LLM to return JSON with:
    { "value": "...", "confidence": 0.0-1.0, "source_snippet": "..." }
    The model should respond with ONLY the JSON.
    """
    label_examples = ", ".join([f"\"{l}\"" for l in label_hints[:6]])
    prompt = (
        "You are a JSON-only extractor. Given the page text, extract the best value for the requested field.\n\n"
        f"Field: {field_key}\n"
        f"Label hints (possible labels or words in the document that indicate this field): {label_examples}\n\n"
        "Rules:\n"
        " - Return EXACTLY one JSON object and nothing else.\n"
        " - JSON schema: {\"value\": \"...\", \"confidence\": 0.0-1.0, \"source_snippet\": \"...\"}\n"
        " - Use \"Not Found\" as value if the field is not present.\n"
        " - Confidence should reflect your certainty (0.0 - 1.0). Use 0.0 for 'Not Found'.\n"
        " - source_snippet should be a short excerpt (<= 120 chars) from the page that supports the value.\n"
        " - Do not hallucinate or invent values.\n\n"
        "Page text:\n"
        f"'''{page_text[:12000]}'''\n\n"
        "Return the JSON now."
    )
    return prompt

def call_gemini_strict(prompt: str, model_choice: str, api_key: str, max_retries: int = 1) -> dict:
    """
    Try to call google.generativeai to extract structured JSON.
    Returns a dict with keys: value, confidence, source_snippet
    If genai not available or call fails, returns {"value": "", "confidence":0.0, "source_snippet": "..."}
    """
    if genai is None:
        return {"value": "", "confidence": 0.0, "source_snippet": "Gemini client not installed"}

    try:
        genai.configure(api_key=api_key)
        # Use chat endpoint if available
        text = ""
        try:
            if hasattr(genai, "chat") and hasattr(genai.chat, "create"):
                resp = genai.chat.create(model=model_choice, messages=[{"role": "user", "content": prompt}])
                # many flavors of response; try to pull best-available text
                # attempt: resp.last.content or resp.choices[0].message.content
                if hasattr(resp, "last"):
                    text = getattr(resp.last, "content", "") or ""
                else:
                    try:
                        text = resp.choices[0].message.content
                    except Exception:
                        text = str(resp)
            elif hasattr(genai, "generate"):
                resp = genai.generate(model=model_choice, prompt=prompt)
                text = getattr(resp, "text", "") or getattr(resp, "output", "") or ""
            else:
                resp = genai.chat.create(model=model_choice, messages=[{"role": "user", "content": prompt}])
                text = str(resp)
        except Exception as e:
            return {"value": "", "confidence": 0.0, "source_snippet": f"LLM call failed: {e}"}

        # Attempt parse JSON
        parsed = None
        try:
            parsed = json.loads(text)
        except Exception:
            # find JSON object inside text
            m = re.search(r"\{.*\}", text, re.S)
            if m:
                try:
                    parsed = json.loads(m.group(0))
                except Exception:
                    parsed = None

        if not parsed or not isinstance(parsed, dict):
            return {"value": "", "confidence": 0.0, "source_snippet": text[:200]}

        # normalize keys
        value = str(parsed.get("value", "")).strip()
        try:
            confidence = float(parsed.get("confidence", 0.0))
            confidence = max(0.0, min(1.0, confidence))
        except Exception:
            confidence = 0.0
        snippet = str(parsed.get("source_snippet", "")).strip()
        return {"value": value if value else "Not Found", "confidence": confidence, "source_snippet": snippet}

    except Exception as e:
        return {"value": "", "confidence": 0.0, "source_snippet": f"Gemini error: {e}"}

# ---------- Aggregation & merge ----------

def merge_page_level_results(per_page_results: List[dict]) -> dict:
    """
    Accepts a list of results (each: {"value", "confidence", "source_snippet"}) across pages.
    Merge into one final value:
    - If any page has confidence >= 0.8, pick the highest confidence page.
    - Else if any non-"Not Found" exists, pick the highest confidence.
    - Else return Not Found with confidence 0.0.
    """
    if not per_page_results:
        return {"value": "Not Found", "confidence": 0.0, "source_snippet": ""}

    # sort by confidence desc
    sorted_pages = sorted(per_page_results, key=lambda x: x.get("confidence", 0.0), reverse=True)
    top = sorted_pages[0]
    if top.get("confidence", 0.0) >= 0.8 and top.get("value", "") != "Not Found":
        return top
    # fallback: find first non-"Not Found"
    for r in sorted_pages:
        if r.get("value") and r.get("value") != "Not Found":
            return r
    # nothing found
    return {"value": "Not Found", "confidence": 0.0, "source_snippet": ""}

# ---------- UI and main flow ----------

st.set_page_config(page_title="RAG Document Filler — improved", layout="wide", page_icon="🧠")
st.title("🧠 RAG-Powered Document Filler — Improved")

with st.sidebar:
    st.header("Settings")
    env_key = os.getenv("GEMINI_API_KEY", "")
    api_key = st.text_input("Gemini API Key (optional)", value=env_key, type="password")
    model_choice = st.selectbox("LLM Model (Gemini)", ["gemini-2.5-flash", "gemini-2.5-pro"], index=0)
    use_llm = st.checkbox("Enable LLM fallback (Gemini) for low-confidence fields", value=bool(api_key))
    max_pages_to_process = st.number_input("Max pages to process per source (0 = all)", min_value=0, max_value=1000, value=0)

st.markdown("Upload a DOCX template with placeholders like `[INSURED_NAME]` and source PDFs/images. Use the confirm table to correct before download.")

col1, col2 = st.columns(2)
with col1:
    template_file = st.file_uploader("Upload template (.docx)", type="docx")
with col2:
    source_files = st.file_uploader("Upload sources (PDF or images)", accept_multiple_files=True, type=["pdf", "png", "jpg", "jpeg", "tiff", "bmp", "heic"])

if st.button("Start Extraction"):
    if not template_file or not source_files:
        st.error("Please upload both template and source files.")
        st.stop()

    # progress bar (0.0 - 1.0)
    progress = st.progress(0.0)
    status = st.empty()

    # Read template and get placeholders
    try:
        template_bytes = template_file.read()
        doc = Document(io.BytesIO(template_bytes))
    except Exception as e:
        st.error(f"Failed to read template: {e}")
        st.stop()

    placeholder_pattern = r'\[([^\]]+)\]'
    placeholders = set()
    for paragraph in doc.paragraphs:
        placeholders.update(re.findall(placeholder_pattern, paragraph.text))
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    placeholders.update(re.findall(placeholder_pattern, paragraph.text))
    placeholders = sorted([p.strip() for p in placeholders if p.strip()])

    if not placeholders:
        st.warning("No placeholders found in template.")
        st.stop()

    status.text(f"Found {len(placeholders)} placeholders.")
    progress.progress(0.05)

    # Extract pages for each uploaded source
    status.text("Extracting text from source files (page-level)...")
    all_pages: List[Tuple[str, str]] = []  # list of (source_filename, page_text)
    processed_pages = 0
    for f in source_files:
        # create temp file
        ext = os.path.splitext(f.name)[1]
        with tempfile.NamedTemporaryFile(delete=False, suffix=ext) as tmp:
            tmp.write(f.getvalue())
            tmp_path = tmp.name

        try:
            if f.name.lower().endswith(".pdf"):
                pages = extract_pages_from_pdf(tmp_path)
                if max_pages_to_process > 0:
                    pages = pages[:max_pages_to_process]
                for p in pages:
                    all_pages.append((f.name, p))
            else:
                text = extract_text_from_image_file(tmp_path)
                all_pages.append((f.name, text))
        finally:
            try:
                os.unlink(tmp_path)
            except Exception:
                pass

    if not all_pages:
        st.warning("No text found in the uploaded sources.")
        progress.progress(0.1)
        st.stop()

    status.text(f"Pages extracted: {len(all_pages)}")
    progress.progress(0.12)

    # Per-field extraction: iterate fields, examine each page chunk, optionally call LLM
    status.text("Extracting fields across pages...")
    per_field_page_results: Dict[str, List[dict]] = {k: [] for k in placeholders}

    # define label map: prefer user-defined mapping, else fallback to DEFAULT_FIELD_LABEL_MAP
    field_label_map = DEFAULT_FIELD_LABEL_MAP  # you can make this editable in UI if needed

    total_fields = len(placeholders)
    for idx, field_key in enumerate(placeholders):
        status.text(f"Extracting: {field_key} ({idx+1}/{total_fields})")
        # build label hints (if known) else use tokens from field name
        label_hints = field_label_map.get(field_key, [])
        if not label_hints:
            # turn underscores into words
            label_hints = [" ".join(part.capitalize() for part in field_key.split("_"))]

        # examine pages (we'll limit the number of pages for LLM calls)
        for page_i, (src_name, page_text) in enumerate(all_pages):
            if not page_text or len(page_text.strip()) < 5:
                per_field_page_results[field_key].append({"value": "Not Found", "confidence": 0.0, "source_snippet": ""})
                continue

            # quick heuristic local extraction: look for label tokens in page_text
            page_lower = page_text.lower()
            heur_found = False
            heur_value = ""
            heur_snip = ""
            heur_conf = 0.0

            # scan for patterns like "Label: value" or "Label - value"
            for lh in label_hints:
                # check a few variants
                patterns = [
                    rf"{re.escape(lh)}\s*[:\-]\s*(.+?)(?:\n|$)",
                    rf"{re.escape(lh)}\s+is\s+(.+?)(?:\n|$)",
                    rf"{re.escape(lh)}\s+(.+?)(?:\n|$)"
                ]
                for pat in patterns:
                    m = re.search(pat, page_text, flags=re.I)
                    if m:
                        heur_found = True
                        heur_value = m.group(1).strip()[:300]
                        heur_snip = heur_value[:120]
                        heur_conf = 0.6  # moderate confidence for heuristic
                        break
                if heur_found:
                    break

            # also try searching for tokens from field name if heuristics didn't find
            if not heur_found:
                tokens = [t.lower() for t in re.findall(r"\w+", field_key) if len(t) > 2]
                found_tokens = [t for t in tokens if t in page_lower]
                if found_tokens:
                    # take the full line containing tokens
                    lines = [ln.strip() for ln in page_text.splitlines() if ln.strip()]
                    line_match = ""
                    for ln in lines:
                        if any(t in ln.lower() for t in found_tokens):
                            line_match = ln
                            break
                    if line_match:
                        heur_found = True
                        heur_value = line_match[:300]
                        heur_snip = heur_value[:120]
                        heur_conf = 0.45

            if heur_found:
                per_field_page_results[field_key].append({"value": heur_value or "Not Found", "confidence": heur_conf, "source_snippet": heur_snip})
            else:
                # no heuristic found; mark Not Found for now
                per_field_page_results[field_key].append({"value": "Not Found", "confidence": 0.0, "source_snippet": ""})

        # optionally use LLM fallback for this field if many pages have low confidence and LLM enabled
        page_results = per_field_page_results[field_key]
        # if we already have a strong result (confidence >=0.8) skip LLM
        if any(r.get("confidence", 0.0) >= 0.8 and r.get("value") != "Not Found" for r in page_results):
            pass
        else:
            # count pages with any non-notfound value
            non_found_count = sum(1 for r in page_results if r.get("value") and r.get("value") != "Not Found")
            low_conf_count = sum(1 for r in page_results if r.get("confidence", 0.0) < 0.5)

            if use_llm and genai is not None and (low_conf_count > 0):
                # call LLM per page, but limit to e.g. first 10 pages to avoid cost/truncation
                pages_for_llm = all_pages[:min(len(all_pages), 10)]
                for src_name, page_text in pages_for_llm:
                    prompt = build_strict_prompt_for_field(field_key, label_hints, page_text)
                    parsed = call_gemini_strict(prompt, model_choice, api_key)
                    per_field_page_results[field_key].append(parsed)
            # else: skip LLM

        # update progress (scale roughly: start 0.12 -> finish 0.6 while field extraction runs)
        progress_val = 0.12 + 0.45 * ((idx + 1) / total_fields)
        progress.progress(min(0.9, progress_val))

    # Now merge page-level results into final extracted_data
    extracted_data: Dict[str, Dict[str, Any]] = {}
    for key in placeholders:
        merged = merge_page_level_results(per_field_page_results.get(key, []))
        extracted_data[key] = merged

    status.text("Extraction complete. Please review & edit any low-confidence fields below.")
    progress.progress(0.92)

    # Show a confirmation editable table: user can override any value before writing to DOCX
    st.subheader("Extracted Fields — Review & Edit")
    rows = []
    for k, v in extracted_data.items():
        rows.append({
            "field": k,
            "value": v.get("value", "Not Found"),
            "confidence": round(float(v.get("confidence", 0.0)), 3),
            "source_snippet": v.get("source_snippet", "")
        })

    import pandas as pd
    df = pd.DataFrame(rows)
    edited_df = st.experimental_data_editor(df, num_rows="fixed", use_container_width=True)

    # When user confirms, replace placeholders and allow download
    if st.button("Finalize & Download Filled Document"):
        # update extracted_data from edited_df
        for _, row in edited_df.iterrows():
            key = row["field"]
            val = row["value"] if row["value"] not in (None, "") else "Not Found"
            extracted_data[key] = {
                "value": val,
                "confidence": float(row.get("confidence", 0.0)),
                "source_snippet": row.get("source_snippet", "") or ""
            }

        # Replace placeholders in doc
        def replace_in_paragraph(paragraph, key, value):
            placeholder = f"[{key}]"
            if placeholder in paragraph.text:
                for run in paragraph.runs:
                    if placeholder in run.text:
                        run.text = run.text.replace(placeholder, str(value))

        try:
            # create a copy of the Document in memory and modify it
            out_doc = Document(io.BytesIO(template_bytes))
            for p in out_doc.paragraphs:
                for key in placeholders:
                    replace_in_paragraph(p, key, extracted_data.get(key, {}).get("value", ""))
            for table in out_doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            for key in placeholders:
                                replace_in_paragraph(p, key, extracted_data.get(key, {}).get("value", ""))

            bio = io.BytesIO()
            out_doc.save(bio)
            bio.seek(0)
        except Exception as e:
            st.error(f"Failed to create output document: {e}")
            st.stop()

        progress.progress(1.0)
        status.text("Done — download your filled document")
        st.download_button(
            "📥 Download filled document",
            data=bio,
            file_name=f"RAG_Filled_{template_file.name.replace('.docx','')}_{datetime.now().strftime('%Y%m%d')}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

    # show diagnostics (optional)
    if st.checkbox("Show raw per-page extraction (diagnostics)"):
        st.write("Per-page results (sample):")
        for k, per_pages in per_field_page_results.items():
            st.write(f"--- {k} ---")
            st.json(per_pages[:8])

    st.success("Pipeline finished (you can re-run with different files or change settings).")
